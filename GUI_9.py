import wx
import os
from SQLite_3 import Database
from Excel_3 import xlrd_fx
from Excel_3 import Macros
from Excel_3 import XL_Constants
import xlrd
import time
import datetime
import sqlite3
from docx import Document
from docx.shared import Inches
from DataVis import DataVisualizer
import win32com.client as win32
from shutil import copyfile

class ManHourTrackerFrame(wx.Frame):
	def __init__(self, parent, title):
		wx.Frame.__init__(self, parent, title=title, size=(560,700))
		
		self.CreateStatusBar() #A Status bar in the bottom of the window
		
		#Set up the menu which will populate the menubar.
		SelectReportStatusBar = 'Defines the *.docx file the queries will be printed to.'
		filemenu = wx.Menu()
		menuSelectReport = filemenu.Append(wx.ID_OPEN, "&Select Report", " Defines the *.docx file the queries will be printed to.")
		menuBackupDatabase = filemenu.Append(wx.ID_ANY, "&Database Backup", " Defines database backup location.")
		menuAbout = filemenu.Append(wx.ID_ABOUT, "&About", " Provides a background about this program.")
		filemenu.AppendSeparator()
		menuExit = filemenu.Append(wx.ID_EXIT,"E&xit", " Exit the program.")
		
		#Create the menubar.
		menuBar = wx.MenuBar()
		menuBar.Append(filemenu,"&File") # Adding the "filemenu" to the MenuBar
		self.SetMenuBar(menuBar) # Adding the MenuBar to the Frame content.
		self.Show(True)
		
		#Bind menu items
		self.Bind(wx.EVT_MENU, self.OnBackupDatabase, menuBackupDatabase)
		self.Bind(wx.EVT_MENU, self.OnSelectReport, menuSelectReport)
		self.Bind(wx.EVT_MENU, self.OnAbout, menuAbout)
		self.Bind(wx.EVT_MENU, self.OnExit, menuExit)
		self.Bind(wx.EVT_CLOSE, self.OnExit)
	
	def OnBackupDatabase(self, event):
		panel.EventLogger('Sorry this feature is under construction.')
		panel.EventLogger('Action Complete\n')
		
	def OnSelectReport(self, event):  
		#Define the report file 
		self.dirname = ''
		dlg = wx.FileDialog(self, "Select an Existing Document or Name a New a Document", self.dirname, "", "*.docx", wx.OPEN)
		if dlg.ShowModal() == wx.ID_OK:
			self.filename = dlg.GetFilename()
			self.dirname = dlg.GetDirectory()
			self.fullpath = dlg.GetPath()
			dlg.Destroy()
			panel.EventLogger('Report output file set: %s' % self.fullpath)
			panel.ReportFileName = self.filename
			panel.ReportFilePath = self.fullpath
			panel.ReportPath = self.dirname
			panel.EventLogger('Action Complete\n')
			
	def OnAbout(self, event):
		message = """
		TIME TRAX 2.2 is a software platform designed to help LECO's
		Service Department managers allocate their resources more efficiently.
		
		System Requirements & Python Libraries:
		\t-Windows 7 64 bit
		\t-Office 2013
		\t-Python 2.7.9 (64 bit)
		\t-wxPython 3.0 win64
		\t-Python 2.7 pywin32-220
		\t-sqlite3
		\t-matplotlib (1.5.1)
		\t-numpy (1.11.1)
		\t-python-docx (0.8.6)
		\t-xlrd (1.0.0)
		\t-lxml (3.6.0)
		"""
		title = "TIME TRAX 2.2: About"
		dlg = wx.MessageDialog(self, message, title, wx.OK)
		dlg.ShowModal() #Show it
		dlg.Destroy() #Finally destroy it when finished.
		
	def OnExit(self, event):
		panel.EventLogger('Saving & Disconnecting from databases.\n\t\tPlease Wait...')
		panel.catDB.conn.commit()
		panel.empDB.conn.commit()
		panel.catDB.conn.close()
		panel.empDB.conn.close()
		
		if panel.real_connect == False:
			try:
				copyfile(panel.empDB_FilePath, panel.backup_empDB_FilePath)
				copyfile(panel.catDB_FilePath, panel.backup_catDB_FilePath)
				print "copied"
			except IOError:
				pass
		else:
			copyfile(panel.dbpath + "\\SepSci_cat.db", panel.ActiveDirectory + "\\SepSci_cat.db")
			copyfile(panel.dbpath + "\\SepSci_emp.db", panel.ActiveDirectory + "\\SepSci_emp.db")
			
		self.Destroy()

		
class ManHourTrackerPanel(wx.Panel):
	
	def __init__(self, parent):
		wx.Panel.__init__(self, parent)
		
		
		self.ReportFileName = ''
		self.ReportFilePath = ''
		self.ReportPath = ''
		self.temp_image_file_created = ''
		self.temp_docx_file_created = ''
		self.data_aggregation = ''
		self.hour_label = ''
		self.average_by_label = ''
		
		self.ActiveDirectory = 'C:\\TimeTrax2.2\\'
		
		self.WordTemplate_file_path = 'C:\\ProgramData\\TimeTrax\\docs\\Template.docx'
		self.ExcelMacros_file_path = 'C:\\ProgramData\\TimeTrax\\docs\\XLManipulation.xlsm'
		self.DB_Backup_Location = 'L:\\SEP SCI Backup Data\\Installation Data\\Pegasus\\3050_UC Berkeley 10-08-2010\\'
		self.dbpath = 'C:\\ProgramData\\TimeTrax\\'
		
		department = "SepSci"
		#department = "GDS"
		
		if department == "SepSci":
			self.category_db_name = 'SepSci_cat.db'
			self.employee_db_name = 'SepSci_emp.db'
			self.Default_Dept_Integer = 0
			self.DepartmentSelected = "SepSci"
		elif department == "GDS":
			self.category_db_name = 'GDS_cat.db'
			self.employee_db_name = 'GDS_emp.db'
			self.Default_Dept_Integer = 1
			self.DepartmentSelected = "GDS"
			
		self.catDB_FilePath = '%s%s' % (self.dbpath, self.category_db_name)
		self.empDB_FilePath = '%s%s' % (self.dbpath, self.employee_db_name)
		
		self.backup_empDB_FilePath = '%s%s' % (self.DB_Backup_Location, self.employee_db_name)
		self.backup_catDB_FilePath = '%s%s' % (self.DB_Backup_Location, self.category_db_name)
		
		# If this is my computer then grab the latest DB copy
		if os.environ['COMPUTERNAME'] == 'CND651145F':
			self.real_connect = True
			try:
				copyfile(self.backup_empDB_FilePath, self.empDB_FilePath)
				copyfile(self.backup_catDB_FilePath, self.catDB_FilePath)
				self.backup_location = True
			except IOError:
				self.backup_location = False
		else:
			self.real_connect = False
		
		'''Section 1 - create sizers'''
		
		#1.)Create main Sizer; all other sizers will be added to this one.
		main_VertSizer = wx.BoxSizer(wx.VERTICAL)
		
		#2.)User input sizer contains the following:
		#query_lists_GridSizer and other_user_input_Vert_Sizer
		user_input_HorzSizer = wx.BoxSizer(wx.HORIZONTAL)
		
		#3.)Query lists GridSizer contains the primary
		#query list box sizers: Population, Category, Start & End dates
		query_lists_GridSizer = wx.GridBagSizer(hgap=1, vgap=1)
		
		#4-7.)Each of the preceding list box sizers consists of 
		#their respective label and the list box
		category_VertSizer = wx.BoxSizer(wx.VERTICAL)
		population_VertSizer = wx.BoxSizer(wx.VERTICAL)
		startdate_VertSizer = wx.BoxSizer(wx.VERTICAL)
		enddate_VertSizer = wx.BoxSizer(wx.VERTICAL)
		
		#8.)Other input options contains: 
		#output options sizer & edit employees grid sizer
		other_user_input_Vert_Sizer = wx.BoxSizer(wx.VERTICAL)
		
		#9.)Output options contains the following:
		#Department Selections radio buttons, Output options radio buttons & Average By list box sizer
		output_options_Vert_Sizer = wx.BoxSizer(wx.VERTICAL)
		
		#10.)Average by list box sizer consists of
		#its label and the list box
		averageby_VertSizer = wx.BoxSizer(wx.VERTICAL)
		
		#11.)Employees grid sizer contains labels and
		#the corresponding text boxes as well as the
		#employee type list box
		edit_employee_list_GridSizer = wx.GridBagSizer(hgap=3, vgap=3)
		
		#12.)Buttons sizer consists of the 4 "action" buttons
		buttons_HorzSizer = wx.BoxSizer(wx.HORIZONTAL)
		
		'''Section 2 - create database connections and query initialization data'''
		
		#Create category Database object
		#If a database doesn't already exist 
		#then add the tables
		#self.catDB_FilePath = '%scat.db' % dbpath
		catDB_FilePath_Exists = os.path.exists(self.catDB_FilePath)
		self.catDB = Database(self.catDB_FilePath)
		if not(catDB_FilePath_Exists):
			self.catDB.Create_Category_Tables()
			
		#Build lists and dictionaries for category list box
		
		#Get category dictionary, list, Sep Sci category index list and GDS category index list 
		self.display_cat_dict = self.catDB.Get_Diplay_Category_Table_Dict()
		self.query_gds_cat_list = self.catDB.Get_GDS_Query_Category_Table_List()
		self.query_ss_cat_list = self.catDB.Get_SS_Query_Category_Table_List()
		self.query_cat_list = self.catDB.Get_Query_Category_Table_List()

		
		#Populate Sep Sci display category list	
		#Inserting 'All', 'Serv./Train./Valid.', 'Road / Home' & 'Sep Sci / GDS' at the beginning of the list
		self.display_sep_sci_cat_list = []
		for item in self.query_ss_cat_list:
			self.display_sep_sci_cat_list.append(self.display_cat_dict[item])
		self.display_sep_sci_cat_list.insert(0,'Sep Sci / GDS')
		self.display_sep_sci_cat_list.insert(0,'Road / Home')
		self.display_sep_sci_cat_list.insert(0,'Serv. / Train. / Valid.')
		self.display_sep_sci_cat_list.insert(0,'All (Excludes Validation)')
		self.display_sep_sci_cat_list.insert(0,'All')
		
		#Populate GDS display category list
		#Inserting 'All', 'Serv./Train./Valid.', 'Road / Home' & 'Sep Sci / GDS' at the beginning of the list		
		self.display_gds_cat_list = []
		for item in self.query_gds_cat_list:
			self.display_gds_cat_list.append(self.display_cat_dict[item])
		self.display_gds_cat_list.insert(0,'Sep Sci / GDS')
		self.display_gds_cat_list.insert(0,'Road / Home')
		self.display_gds_cat_list.insert(0,'Serv. / Train.')
		self.display_gds_cat_list.insert(0,'All')
		
		
		#Create employee Database object
		#If a database doesn't already exist 
		#then add the tables
		#self.empDB_FilePath = '%semp.db' % dbpath
		empDB_FilePath_Exists = os.path.exists(self.empDB_FilePath)
		self.empDB = Database(self.empDB_FilePath)
		if not(empDB_FilePath_Exists):
			self.empDB.Create_Employee_Tables()
			self.empDB.Update_Employee_Tables(department)
			
		#Get the employee names from the employee tables
		#and put them into a list.  Then insert into the list
		#'all', 'field service only', 'in-house only'
		emp_names_DB_Object = self.empDB.Select_Query('employee_name', 'employee', 'ORDER BY employee_name ASC')
		self.emp_names_list = []
		for name_tuple in emp_names_DB_Object:
			name_str = "%s" % name_tuple
			self.emp_names_list.append(str(name_str))
		self.emp_names_list.insert(0,'Field Service')
		self.emp_names_list.insert(0,'In-House')
		self.emp_names_list.insert(0,'Each Individual')
		self.emp_names_list.insert(0,'All')
		
		#Get all the distinct dates from the catDB.
		#Since each table in the catDB has the same
		#data_dates the Miscellaneous table was 
		#arbitrarily selected to query

		data_date_lst = self.Get_Data_Dates()
		
		'''Section 3 - Add widgets to query_lists_GridSizer (3-7)'''
		
		#build category list box & label #4
		if department == "SepSci":
			temp_list =  self.display_sep_sci_cat_list
		elif department == "GDS":
			temp_list = self.display_gds_cat_list
		
		self.lbl_cat_listbox = wx.StaticText(self, label=" Category:")
		self.cat_listbox = wx.ListBox(self, 26, wx.DefaultPosition, (170, 130), temp_list, wx.LB_SINGLE)
		self.cat_listbox.SetSelection(0)
		category_VertSizer.Add(self.lbl_cat_listbox, wx.ALIGN_LEFT)
		category_VertSizer.Add(self.cat_listbox, wx.ALIGN_LEFT)
		
		#build employee list box #5
		self.lbl_emp_listbox = wx.StaticText(self, label=" Population:")
		self.emp_listbox = wx.ListBox(self, 26, wx.DefaultPosition, (170, 130), self.emp_names_list, wx.LB_SINGLE)
		self.emp_listbox.SetSelection(0)
		population_VertSizer.Add(self.lbl_emp_listbox, wx.ALIGN_LEFT)
		population_VertSizer.Add(self.emp_listbox, wx.ALIGN_LEFT)
		
		#Selections for #6 & #7
		if data_date_lst[0] != 'No Data':
			last_date = len(data_date_lst)-1
			last_date_minus_a_week = len(data_date_lst)-7
		else:
			last_date = 0
			last_date_minus_a_week = 0
		
		#build start date list box #6
		self.lbl_start_date_listbox = wx.StaticText(self, label=" Start Date:")
		self.start_date_listbox = wx.ListBox(self, 26, wx.DefaultPosition, (170, 130), data_date_lst, wx.LB_SINGLE)
		self.start_date_listbox.SetSelection(last_date_minus_a_week)
		startdate_VertSizer.Add(self.lbl_start_date_listbox, wx.ALIGN_LEFT)
		startdate_VertSizer.Add(self.start_date_listbox, wx.ALIGN_LEFT)
		
		#build end date list box #7
		self.lbl_end_date_listbox = wx.StaticText(self, label=" End Date:")
		self.end_date_listbox = wx.ListBox(self, 26, wx.DefaultPosition, (170, 130), data_date_lst, wx.LB_SINGLE)
		self.end_date_listbox.SetSelection(last_date)
		enddate_VertSizer.Add(self.lbl_end_date_listbox, wx.ALIGN_LEFT)
		enddate_VertSizer.Add(self.end_date_listbox, wx.ALIGN_LEFT)
		
		#Add list box sizers to the query_lists_GridSizer #3
		query_lists_GridSizer.Add(population_VertSizer, pos=(0,0))
		query_lists_GridSizer.Add(category_VertSizer, pos=(0,1))
		query_lists_GridSizer.Add(startdate_VertSizer, pos=(1,0))
		query_lists_GridSizer.Add(enddate_VertSizer, pos=(1,1))
		
		'''Section 4 - Add widgets to output_options_Vert_Sizer (9-11)'''
		
		#build the Department Selections radio button widget
		radioList = ['Sep. Sci.', 'GDS']
		self.department_radiobut = wx.RadioBox(self, label="Department Categories",  
			choices=radioList, style=wx.RA_SPECIFY_COLS)
		self.department_radiobut.SetSelection(self.Default_Dept_Integer)
		output_options_Vert_Sizer.Add(self.department_radiobut)
		self.Bind(wx.EVT_RADIOBOX, self.EvtDepartmentRadioBox, self.department_radiobut)
		output_options_Vert_Sizer.Add((20,13))
		
		#build the Output Options radio button widget
		radioList = ['Summation', 'Average']
		self.output_options_radiobut = wx.RadioBox(self, label="Output Options",  
			choices=radioList, style=wx.RA_SPECIFY_COLS)
		output_options_Vert_Sizer.Add(self.output_options_radiobut)
		self.Bind(wx.EVT_RADIOBOX, self.EvtOutputOptionsRadioBox, self.output_options_radiobut)
		
		#build the average by list box #10
		ave_by_List = ['Day', 'Week', 'Month', 'Quarter', 'Year']
		self.lbl_ave_by_listbox = wx.StaticText(self, label="  Average By:")
		self.ave_by_listbox = wx.ListBox(self, 26, wx.DefaultPosition, (182, 45), ave_by_List, wx.LB_SINGLE)
		self.ave_by_listbox.SetSelection(1)
		self.ave_by_listbox.Disable()
		averageby_VertSizer.Add(self.lbl_ave_by_listbox, wx.ALIGN_LEFT)
		averageby_VertSizer.Add(self.ave_by_listbox, wx.ALIGN_LEFT)
		
		#build clock number widget
		self.lbl_clock_no = wx.StaticText(self, label=" Clock No: ")
		self.clock_no_txtCtrl = wx.TextCtrl(self, size=(122,-1))
		
		#build employee name widget
		self.lbl_employee_name = wx.StaticText(self, label=" Name: ")
		self.employee_name_txtCtrl = wx.TextCtrl(self, size=(122,-1))
		
		#build employee type widget
		emp_type_List = ['Field', 'In-House', 'N/A']
		self.lbl_type = wx.StaticText(self, label=" Type: ")
		self.type_listbox = wx.ListBox(self, 26, wx.DefaultPosition, (122,-1), emp_type_List, wx.LB_SINGLE)
		self.type_listbox.SetSelection(0)
		
		#populate edit_employee_list_GridSizer #11
		edit_employee_list_GridSizer.Add(self.lbl_clock_no, pos=(0,0))
		edit_employee_list_GridSizer.Add(self.clock_no_txtCtrl, pos=(0,1))
		edit_employee_list_GridSizer.Add(self.lbl_employee_name, pos=(1,0))
		edit_employee_list_GridSizer.Add(self.employee_name_txtCtrl, pos=(1,1))
		edit_employee_list_GridSizer.Add(self.lbl_type, pos=(2,0))
		edit_employee_list_GridSizer.Add(self.type_listbox, pos=(2,1))
		
		#populate output_options_Vert_Sizer #9
		output_options_Vert_Sizer.Add(averageby_VertSizer)
		output_options_Vert_Sizer.Add((20,18))
		output_options_Vert_Sizer.Add(edit_employee_list_GridSizer)
		
		'''Section - 5 Add widgets to buttons_HorzSizer'''
		
		self.import_btn = wx.Button(self, label = "Import")
		self.report_btn = wx.Button(self, label = "Append Report")
		self.employee_info_btn = wx.Button(self, label = "Employee Info")
		self.update_employee_btn = wx.Button(self, label = "Update Employee")
		
		buttons_HorzSizer.Add(self.import_btn)
		buttons_HorzSizer.Add(self.report_btn)
		buttons_HorzSizer.Add(self.employee_info_btn)
		buttons_HorzSizer.Add(self.update_employee_btn)
		
		self.Bind(wx.EVT_BUTTON, self.OnImport, self.import_btn)
		self.Bind(wx.EVT_BUTTON, self.OnReport, self.report_btn)
		self.Bind(wx.EVT_BUTTON, self.OnUpdateEmployee, self.update_employee_btn)
		self.Bind(wx.EVT_BUTTON, self.OnEmployeeInfo, self.employee_info_btn)
		
		'''Section 6 - Add all sizers to the main_VertSizer'''
		user_input_HorzSizer.Add(query_lists_GridSizer, wx.ALIGN_TOP)
		user_input_HorzSizer.Add(output_options_Vert_Sizer, wx.ALIGN_TOP)
		main_VertSizer.Add(user_input_HorzSizer, wx.ALIGN_LEFT)
		
		#create status output
		main_VertSizer.Add((20,2))
		self.lbl_status_logger = wx.StaticText(self, label=" Status Output: ")
		main_VertSizer.Add(self.lbl_status_logger, wx.ALIGN_LEFT)
		self.status_logger = wx.TextCtrl(self, size=(523,175), style=wx.TE_MULTILINE | wx.TE_READONLY)
		main_VertSizer.Add(self.status_logger, wx.ALIGN_LEFT)
		
		#Create initial status message
		current_time = str(time.strftime("%H:%M:%S"))
		current_date = str(time.strftime("%m/%d/%Y"))
		message = "Report File Not Set\n"
		initial_status_message = "%s %s - %s\n" % (current_date, current_time, message) 
		self.status_logger.AppendText(initial_status_message)

		if self.real_connect == True and self.backup_location == True:
			initial_status_message = "\tReal Connect = %s\n\tDB Transfer Successful\n" % str(self.real_connect)
			self.status_logger.AppendText(initial_status_message)

		
		#add button sizer
		main_VertSizer.Add(buttons_HorzSizer)
		self.SetSizerAndFit(main_VertSizer)
	
	def OnImport(self, event):
		#self.catDB
		#self.empDB
		self.dirname = ''
		dlg = wx.FileDialog(self, "Choose a file", self.dirname, "", "*.xlsx", wx.FD_MULTIPLE)
		if dlg.ShowModal() == wx.ID_OK:
				
			#get full excel paths and file names from dialogue objects
			xl_fullpaths_lst = dlg.GetPaths()
			xl_filenames_lst = dlg.GetFilenames() 
			
			#Create an XL_Constants object
			#get the row and column locations
			xl_row_columns = XL_Constants()
			
			clk_no_row = xl_row_columns.clk_no_row
			monday_row = xl_row_columns.monday_row
			clk_no_and_monday_column = xl_row_columns.clk_no_and_monday_column
			day_date_row = xl_row_columns.day_date_row
			start_column_D = xl_row_columns.start_column_D
			version_column = xl_row_columns.version_column
			version_row = xl_row_columns.version_row
			
			category_start_row = xl_row_columns.Get_Category_Start_Row_Dict()
				
			#Create an excel macros object
			excel_macros = Macros()
				
			#iterate through excel files
			for xl_index, xl_fullpaths in enumerate(xl_fullpaths_lst):
			
				#Print to status the full file path of the ith excel file
				#self.EventLogger(str(xl_fullpaths))
				
				#create xlsx_file object from defined xlsx file path
				#if xlrd cannot open the file then the excel file was last
				#saved using a version older than 2013
				#To correct open the file using win32com.client and
				#run a macro the opens the the excel in 2013 then save it.
				#Then xlrd will be able to open the file.
				try:
					xlsx = xlrd_fx(str(xl_fullpaths),0)
				except xlrd.biffh.XLRDError:
					#create an old_file object which opens the excel file containing
					#the macro which opens, saves, and closes given excel file
					excel_macros.OldFileUpgrade(xl_fullpaths)
					xlsx = xlrd_fx(str(xl_fullpaths),0)
					
				#Get version number & run checks
				
				version_no = str(xlsx.get_cell_value(version_row, version_column, "Not Activity Sheet"))
				valid_activity_sheet_version = False
				
				if self.clock_no_txtCtrl.GetValue() == 'kite' and self.employee_name_txtCtrl.GetValue() == 'convert':
					if version_no == 'v1.5':
						message = "Version v1.5 cannot be upgraded - Upgrade Skipped"
					elif version_no == 'v2.0':
						message = "File is already v2.0 - Upgrade Skipped"
					elif version_no == '!^$dfuj4862':
						excel_macros.TransferToV2_FromOriginal(xl_fullpaths)
						message = "File Successfully Converted to v2.0"
					else:
						excel_macros.TransferToV2(xl_fullpaths)
						message = "File Successfully Converted to v2.0"
				elif version_no == "Not Activity Sheet":
					message = 'Excel file imported is not an activity sheet' 
				elif version_no == '!^$dfuj4862':
					message = 'Activity sheet version 1.0 is no longer supported'
				elif version_no == '#*26554MHJ~':
					message = 'Activity sheet version 1.1 is no longer supported'
				elif version_no != 'v2.0':
					message = 'Activity sheet version %s is no longer supported' % version_no
				elif version_no == 'v2.0':
					valid_activity_sheet_version = True
					message = "File is ok"
				else:
					message = "Invalid excel file."
					
				message += " - %s" % xl_filenames_lst[xl_index]
					
				if valid_activity_sheet_version:
					
					#Get the clock number, then verify that it has an employee match
					
					clk_no = str(xlsx.get_cell_value(clk_no_row, clk_no_and_monday_column, None))
					cursor = self.empDB.conn.execute("SELECT employee_name FROM employee WHERE clk_no = ?", (clk_no,))
					x = cursor.fetchall()
					if x == []:
						valid_clock_number = False
					else:
						valid_clock_number = True
					
					#Get Monday's date, then verify that it is a Monday
					
					monday_date = xlsx.get_cell_value(monday_row, clk_no_and_monday_column, "No Date Entered")
					if not(monday_date == "No Date Entered"):
						monday_date = xlsx.convert_excel_date(monday_date, 'yyyy-mm-dd')
						day_of_week = str(datetime.datetime.strptime(monday_date, '%Y-%m-%d').strftime('%A'))
		
						if day_of_week == 'Monday':
							valid_date = True
						else:
							valid_date = False
							monday_date = str(datetime.datetime.strptime(monday_date, '%Y-%m-%d').strftime('%m/%d/%Y'))
							message = "Monday's Date of %s is invalid. The provided date is a %s." % (monday_date, day_of_week)
					elif monday_date == "No Date Entered":
						valid_date = False
						message = "Monday's Date value is invalid because no value was entered"
						
					#if either clock number or date is not valid then
					#mark the problem field(s), log the event, then skip the file
					
					if valid_clock_number == False or valid_date == False:
						if valid_clock_number == False:
							self.EventLogger("Clock number of %s does not match any of the employees in the database. - %s" % (clk_no, xl_filenames_lst[xl_index]))
						if valid_date == False:
							message += " - %s" % xl_filenames_lst[xl_index]
							self.EventLogger(message)
							
						message = 'ERROR - File not loaded - Invalid value(s) marked within the excel file.\n   %s\n' % xl_fullpaths
						self.EventLogger(message)
						excel_macros.BadValueMarker(xl_fullpaths, valid_date, valid_clock_number)
					
					
					elif valid_clock_number == True and valid_date == True:
						
						#Now that activity file has been validated, check for potential data overwrites
						#SELECT Query upload_dates for mondays_date and given employee_clk_no
						
						cursor = self.empDB.conn.execute("SELECT employee_clk_no mondays_date FROM upload_dates WHERE mondays_date = ? AND employee_clk_no = ?", (monday_date, clk_no,))
						
						x = cursor.fetchall()

						if x == []:
							data_conflict = False
						else:
							data_conflict = True
							#yes = True and no = False
							question = '''There are already entries within the database
										matching the dates found in this activity sheet:
										
										%s
										
										Do you wish to replace the existing database records 
										with information found within this activity sheet?''' % xl_filenames_lst[xl_index]
							caption = "OVERWRITE WARNING - %s" % xl_filenames_lst[xl_index]
							dlg = wx.MessageDialog(self, question, caption, wx.YES_NO | wx.ICON_QUESTION)
							overwrite_data = dlg.ShowModal() == wx.ID_YES
							dlg.Destroy()
							message = "%s - User overwrite data = %s - %s" % ("OVERWRITE WARNING", str(overwrite_data), xl_filenames_lst[xl_index])
							self.EventLogger(message)
							
							if overwrite_data == True:
								#Delete Query to remove conflict data from upload dates table
								
								self.empDB.conn.execute("DELETE FROM upload_dates WHERE employee_clk_no = ? AND mondays_date = ?", (clk_no, monday_date,))
								self.empDB.conn.commit()
								data_conflict = False
								
								#Delete Query to remove conflict data from each of the category tables
								for table in self.query_cat_list:
									#Monday's date is found within the activity sheet; determine Sunday's date
									#to complete the DELETE FROMm date range condition
									sunday_date = self.AddDays(monday_date, 6)
									delete_query_statement = "DELETE FROM %s WHERE employee_clk_no = '%s' AND data_date BETWEEN '%s' AND '%s' \
									" % (table, clk_no, monday_date, sunday_date)
									self.catDB.conn.execute(delete_query_statement)
									self.catDB.conn.commit()
							
						if data_conflict == False:
						
							#Insert Query - update upload_dates table for the given employee
							self.empDB.Insert_Query_No_Conditions('upload_dates', ['employee_clk_no','mondays_date'], [clk_no, monday_date])
							
							#iterate through each of the tables/categories
							for table in self.query_cat_list:
							
								#get the columns / sub-categories
								columns = self.catDB.Get_Columns(table)
								
								#Get the starting row number for the given category from the cat_start_row dictionary
								start_row = category_start_row[table]
								
								#iterate through each day
								for day in range(0,7):
								
									#reset list
									DB_record_list = []
									
									#define column number 
									#Day 0 starts at "D" or 3
									column = start_column_D + day
									
									#Get date of given data day
									Day_Date_Cell = xlsx.get_cell_value(day_date_row, column, None)
									data_date = xlsx.convert_excel_date(Day_Date_Cell, 'yyyy-mm-dd')
									
									#Add the data date and clock number to the DB_record_list 
									DB_record_list = [data_date, clk_no]
									
									#Iterate through each sub-category cell value
									for index_column, table_column in enumerate(columns):
									
										#skip the first two sub-categories (data_date and clock number) 
										#as they do not have corresponding values within these excel tables
										#further more we've already grabbed these values and appended them
										#to the DB_record_list
										
										if index_column > 1:
											#The - 2 accounts for the skipped sub-categories
											row = start_row + index_column - 2
											DB_record_list.append(xlsx.get_cell_value(row, column, 0.00))
									
									#INSERT DB_record_list INTO the corresponding table
									self.catDB.Insert_Query_No_Conditions(table, columns, DB_record_list)
							
							self.EventLogger('Successfully Imported - %s' %xl_filenames_lst[xl_index])
					
				elif not(valid_activity_sheet_version):
					self.EventLogger(message)
	
		self.Update_Date_ListBoxes()
		self.EventLogger('Action Complete\n')
		dlg.Destroy()
		
	def OnReport(self, event):
		#Initial user input checks:
			#1.) Start date <= End date
			#2.) self.ReportFileName != '' (Has the report path been set)
			#3.) self.ReportPath exists (is the file save directory valid)
			#4.) self.ReportFilePath exists (is this a new report or appending an existing one?)
			#5.) if file exists, is the Report file word document currently open if yes then close it(docx throws an error when connecting to an open document)
			
		#Check 1.) Start date <= End date
		
		#Get the start day in mm/dd/yyyy format & yyyy-mm-dd format
		lst_index = self.start_date_listbox.GetSelection()
		start_date_mm_dd_yyyy = str("%s" % self.start_date_listbox.GetString(lst_index))
		start_date_yyyy_mm_dd = str(datetime.datetime.strptime(start_date_mm_dd_yyyy, '%m/%d/%Y  (%a)').strftime("'%Y-%m-%d'"))
		
		#Get the end day in mm/dd/yyyy format & yyyy-mm-dd format
		lst_index = self.end_date_listbox.GetSelection()
		end_date_mm_dd_yyyy = str("%s" % self.end_date_listbox.GetString(lst_index))
		end_date_yyyy_mm_dd = str(datetime.datetime.strptime(end_date_mm_dd_yyyy, '%m/%d/%Y  (%a)').strftime("'%Y-%m-%d'"))
		
		if start_date_yyyy_mm_dd > end_date_yyyy_mm_dd:
			valid_date_range = False
			message = "Invalid Query Condition - Defined Start Date is after the End Date"
			self.EventLogger(message)
		else:
			valid_date_range = True
			
			
		#Check 2.) self.ReportFileName != '' (Has the report path been set)
		if self.ReportFileName != '':
			report_set = True
		else:
			report_set = False
			message = "Invalid Query Condition - The report has not been set."
			self.EventLogger(message)
		
		if report_set == True:
			#Check 3.) self.ReportPath exists (is the file save directory valid)
			valid_path = (os.path.isdir(self.ReportPath))
			if valid_path == False:
				message = "Invalid Query Condition - The report directory is invalid."
				self.EventLogger(message)
			
			#Check 4.) self.ReportFilePath exists (is this a new report or appending an existing one?)
			new_report = not((os.path.exists(self.ReportFilePath)))
			
			#Check 5.) if file exists, is the Report file word document currently open.  
			#If its open then it is closed (p-docx throws an error when working with an open document)
			if new_report == False:
				report_doc = Document(self.ReportFilePath)
				try:
					report_doc.save(self.ReportFilePath)

				except IOError:
					self.close_word_document(True)
			
		if valid_date_range == True and report_set == True and valid_path == True:
			#All checks passed
			
			#Define chart x-axis label and table label based upon self.output_options_radiobut user input
			x = self.output_options_radiobut.GetSelection()
			self.data_aggregation = self.output_options_radiobut.GetString(x)
			if self.data_aggregation == 'Summation':
				self.hour_label = 'Total Hours'
				self.average_by_label = ''
				
			elif self.data_aggregation == 'Average':
				#If average is selected define corresponding title
				#And getting the denominator in the average calculation
				self.hour_label = 'Average Hours'
				average_by_index = self.ave_by_listbox.GetSelection()
				self.average_by_label = str(" per %s" % self.ave_by_listbox.GetString(average_by_index))
				
			#Get all the clock number corresponding to the self.emp_listbox selection
			
			population_index = self.emp_listbox.GetSelection()
			population_selection = str("%s" % self.emp_listbox.GetString(population_index))
			
			if population_index == 0 or population_index == 1:
				#get all the clock numbers in the employee table
				cursor = self.empDB.conn.execute("SELECT clk_no FROM employee ORDER BY employee_name ASC")
				clk_no_list = cursor.fetchall()
				
			if population_index == 2:
				#get all the clock numbers of employees with type = 'In-House' in the employee table
				cursor = self.empDB.conn.execute("SELECT clk_no FROM employee WHERE employee_type = 'In-House' ORDER BY employee_name ASC")
				clk_no_list = cursor.fetchall()
				
			if population_index == 3:
				#get all the clock numbers of employees with type = 'Field' in the employee table
				cursor = self.empDB.conn.execute("SELECT clk_no FROM employee WHERE employee_type = 'Field' ORDER BY employee_name ASC")
				clk_no_list = cursor.fetchall()
				
			if population_index > 3:
				#get the clock number of selected employee in the employee table
				cursor = self.empDB.conn.execute("SELECT clk_no FROM employee WHERE employee_name = ?", (population_selection,))
				clk_no_list = cursor.fetchall()
			
			#if "each individual" was selected for the population the loop_all_Employees list
			#length will loop the through the report generation logic once for each
			#employee, otherwise a place holder value will be used to allow the loop to 
			#be executed a single time
			if population_index != 1:
				loop_all_Employees = False
				Employee_to_iterate = ['loop once']
			else:
				loop_all_Employees = True
				Employee_to_iterate = [x for x in clk_no_list]
			
			#Determine which tables are included in the query corresponding to the self.cat_listbox selection
			category_index = self.cat_listbox.GetSelection()
			category_selection = str("%s" % self.cat_listbox.GetString(category_index))
			table_list = []
			table_aggregation_list = []
			data_label_list = []
			Home_Road_Report = False
			
			if self.DepartmentSelected == "SepSci":
			
				if category_index == 0:
					#All sep sci tables
					for item in self.query_ss_cat_list:
						table_list.append(item)
					table_aggregation_list = [table_list]
						
				elif category_index == 1:
					#All sep sci tables excluding validation
					for index, item in enumerate(self.query_ss_cat_list):
						if index < len(self.query_ss_cat_list)-1:
							table_list.append(item)
					
					table_aggregation_list = [table_list]
							
				elif category_index == 2:
					#Service tables / training tables / validation table
					service_tables = ['SS_Installations','SS_PM_Site_Visits','SS_Rpr_Maint_Site_Visits','SS_Rmt_Hrdwr_Spt',	'SS_Rpr_Mant_RFB_in_House','SS_Inter_Dep_Spt']
					training_tables = ['SS_Rmt_Sftwr_Spt','SS_Doc_Gen','SS_Online_Training','SS_Onsite_Training','SS_In_House_Training']		
					validation_tables = ['Validation_Duties']
					misc_table = ['Miscellaneous']
					table_aggregation_list = [service_tables, training_tables, validation_tables, misc_table]
					data_label_list = ['Service', 'Training', 'Validation', 'Miscellaneous']
					
					
				elif category_index == 3:
					#Home vs Road tables
					Home_Road_Report = True
					home_tables = [x for x in self.catDB.Get_SS_Query_Category_Table_List()]
					home_tables.remove('Miscellaneous')
					road_tables = ['SS_Installations','SS_PM_Site_Visits','SS_Rpr_Maint_Site_Visits']
					misc_table = ['Miscellaneous']
					
					home_categories = [['Installations_Prep', 'Documentation_Admin'], 
										['PM_Prep', 'Documentation_Admin'],
										['Site_Visit_Prep', 'Documentation_Admin'],
										['Rmt_Hardware_Support_day_sum'],['Rmt_Software_Support_day_sum'],['Inst_Repair_Maint_Rfb_In_House_day_sum'],
										['Document_Generation_day_sum'],['Inter_Dep_Spt_day_sum'],['Online_Training_day_sum'],['Onsite_Training_day_sum'],
										['In_House_Training_day_sum'],['Validation_Duties_day_sum']]
										
					road_categories = [['Travel', 'Hands_On_Instrument', 'Troubleshooting', 'Training_by_Engineer', 'Other'],
										['Travel', 'Hands_On_Instrument', 'Troubleshooting', 'Training_by_Engineer', 'Other'],
										['Travel', 'Hands_On_Instrument', 'Troubleshooting', 'Training_by_Engineer', 'Other']]
										
					misc_categories = [['Miscellaneous_day_sum']]
					
					table_aggregation_list = [home_tables, road_tables, misc_table]
					column_aggregation_list = [home_categories, road_categories, misc_categories]
					data_label_list = ['Home', 'Road', 'Miscellaneous']
					
				elif category_index == 4:
					#Sep Sci tables vs GDS tables
					Sep_Sci_tables = [x for x in self.catDB.Get_SS_Query_Category_Table_List()]
					Sep_Sci_tables.remove('Miscellaneous')
					Sep_Sci_tables.remove('Validation_Duties')
					GDS_tables = [x for x in self.catDB.Get_GDS_Query_Category_Table_List()]
					GDS_tables.remove('Miscellaneous')
					misc_table = ['Miscellaneous']
					table_aggregation_list = [Sep_Sci_tables, GDS_tables, misc_table]
					data_label_list = ['Sep Sci', 'GDS', 'Miscellaneous']
					
				elif category_index > 4:
					#Individual Sep Sci tables (categories)
					table_list.append(self.query_ss_cat_list[category_index-5])
					table_aggregation_list = [table_list]
					
			elif self.DepartmentSelected == "GDS":
			
				if category_index == 0:
					#All GDS tables
					for item in self.query_gds_cat_list:
						table_list.append(item)
					table_aggregation_list = [table_list]
					
				elif category_index == 1:
					#Service tables / training tables 
					service_tables = ['GDS_Installations','GDS_PM_Site_Visits','GDS_Rpr_Maint_Site_Visits','GDS_Rmt_Hrdwr_Spt',	'GDS_Rpr_Mant_RFB_in_House','GDS_Inter_Dep_Spt']
					training_tables = ['GDS_Rmt_Sftwr_Spt','GDS_Doc_Gen','GDS_Online_Training','GDS_Onsite_Training','GDS_In_House_Training']
					misc_table = ['Miscellaneous']
					table_aggregation_list = [service_tables, training_tables, misc_table]
					data_label_list = ['Service', 'Training', 'Miscellaneous']
					
				elif category_index == 2:
					#Home vs Road tables
					Home_Road_Report = True
					home_tables = [x for x in self.catDB.Get_GDS_Query_Category_Table_List()]
					home_tables.remove('Miscellaneous')
					road_tables = ['GDS_Installations','GDS_PM_Site_Visits','GDS_Rpr_Maint_Site_Visits']
					misc_table = ['Miscellaneous']
					
					home_categories = [['Installations_Prep', 'Documentation_Admin'], 
										['PM_Prep', 'Documentation_Admin'],
										['Site_Visit_Prep', 'Documentation_Admin'],
										['Rmt_Hardware_Support_day_sum'],['Rmt_Software_Support_day_sum'],['Inst_Repair_Maint_Rfb_In_House_day_sum'],
										['Document_Generation_day_sum'],['Inter_Dep_Spt_day_sum'],['Online_Training_day_sum'],['Onsite_Training_day_sum'],
										['In_House_Training_day_sum']]
										
					road_categories = [['Travel', 'Hands_On_Instrument', 'Troubleshooting', 'Training_by_Engineer', 'Other'],
										['Travel', 'Hands_On_Instrument', 'Troubleshooting', 'Training_by_Engineer', 'Other'],
										['Travel', 'Hands_On_Instrument', 'Troubleshooting', 'Training_by_Engineer', 'Other']]
										
					misc_categories = [['Miscellaneous_day_sum']]
					
					table_aggregation_list = [home_tables, road_tables, misc_table]
					column_aggregation_list = [home_categories, road_categories, misc_categories]
					data_label_list = ['Home', 'Road', 'Miscellaneous']
					
				elif category_index == 3:
					#Sep Sci tables vs GDS tables
					Sep_Sci_tables = [x for x in self.catDB.Get_SS_Query_Category_Table_List()]
					Sep_Sci_tables.remove('Miscellaneous')
					Sep_Sci_tables.remove('Validation_Duties')
					GDS_tables = [x for x in self.catDB.Get_GDS_Query_Category_Table_List()]
					GDS_tables.remove('Miscellaneous')
					misc_table = ['Miscellaneous']
					table_aggregation_list = [GDS_tables, Sep_Sci_tables, misc_table]
					data_label_list = ['GDS', 'Sep Sci', 'Miscellaneous']
					
				elif category_index > 3:
					#Individual GDS tables (categories)
					table_list.append(self.query_gds_cat_list[category_index-4])
					table_aggregation_list = [table_list]
			
			'''Data base quering loop'''
			#Primary function is to populate the data_value_list & data_label_list
			#So that the bar chart and table in word can be populated
			
			for each_employee_clk_no in Employee_to_iterate:
				
				if loop_all_Employees == True:
					#update new_report status
					new_report = not((os.path.exists(self.ReportFilePath)))
					
					#query name of current employee for the "each individual" query request
					cursor = self.empDB.conn.execute("SELECT employee_name FROM employee WHERE clk_no = '%s'" % each_employee_clk_no)
					query_emp_name = cursor.fetchall()
					population_selection = "%s" % query_emp_name[0]
					
					#transfer clock number value to the clk_no_list
					clk_no_list = [each_employee_clk_no]
				
				
				#define / reset data value list
				data_value_list = []
				
				#If query is NOT a table aggregation then define / reset data value list
				if len(table_aggregation_list) == 1:
					data_label_list = []
				
				for tables_index, tables in enumerate(table_aggregation_list):
				
					#If query IS a table aggregation then define / reset hours
					if len(table_aggregation_list) != 1:
						hours = 0.00
				
					for table_index, table in enumerate(tables):
					
						#define / reset columns_list
						columns_list = []
					
						#Get column(s) from defined table
						all_table_columns_list = list(self.catDB.Get_Columns(table))
						
						#Get columns specific to the defined query
						if Home_Road_Report == True:
							columns_list = column_aggregation_list[tables_index][table_index]
						elif len(tables) == 1:
							#A single table (category) has been selected
							columns_list = all_table_columns_list[3:]
						else: 
							columns_list.append(all_table_columns_list[2])

						for column in columns_list:

							#If query is NOT a table aggregation then define / reset hours
							if len(table_aggregation_list) == 1:
								hours = 0.00
							
							#set the data label depending if a single or all categories have been selected
							if len(columns_list) == 1:
								label = self.display_cat_dict[table]
							else:
								label = column.replace('_',' ')
							
							#Iterate through clk_no_list and query based upon user entered conditions
							for clk_no in clk_no_list:
								clk = str("'%s'" % clk_no)
								
								select_query_statement = "SELECT SUM(%s) FROM %s WHERE employee_clk_no = %s AND data_date BETWEEN %s AND %s \
									" % (column, table, clk, start_date_yyyy_mm_dd, end_date_yyyy_mm_dd)
								
								cursor = self.catDB.conn.execute(select_query_statement)
								result = cursor.fetchall()
								temp_value = str("%s" % result[0])
								
								#if there are no data dates for a given employee for a date 
								#range then none will be returned.  The following conditional
								#statement prevents the ValueError that would result from the
								#attempted float conversion below
								if temp_value == 'None':
									temp_value = 0.00
								hours += float(temp_value)
							
							#If query is NOT a table aggregation then append data lists
							if len(table_aggregation_list) == 1:
								data_value_list.append(hours)
								data_label_list.append(label)
					
					#If query IS a table aggregation ("Serv. / Train. / Valid.", "Road / Home", "Sep / GDS") query then append data lists
					if len(table_aggregation_list) != 1:
						data_value_list.append(hours)
			
				#If the user has selected to average the data
				#execute the corresponding divisor depending on the average-by selection
				if self.data_aggregation == 'Average':
					denominator = self.get_averaging_denominator(average_by_index, start_date_yyyy_mm_dd, end_date_yyyy_mm_dd)
					data_value_list[:] = [round(val / denominator,1) for val in data_value_list]

				#Create bar chart, in doing so return the following:
					#data_value_list as a list of percentages
					#Boolean if the chart_created event was successful
					#Message describing success or failure of chart create event which will be the report query message prefix if 
						#the chart_created event was successful
				percentages, chart_created, chart_message = self.CreateBarChart(data_label_list, data_value_list, columns_list)
				
				#Report query message suffix - it lists all the the query conditions
				message = " - Query Conditions:\n\tPopulation: %s\n\tCategory: %s\n\tStart Date: %s\n\tEnd Date: %s\n\tDepartment: %s\n\tData Aggregation: %s%s" % (population_selection, category_selection, start_date_mm_dd_yyyy, end_date_mm_dd_yyyy, self.DepartmentSelected, self.hour_label, self.average_by_label)
				
				#If a chart is created then the report is appended to the user defined document
				if chart_created == True:
					report_success = self.AppendDocument(new_report, data_label_list, data_value_list, population_selection, percentages, category_selection, start_date_mm_dd_yyyy, end_date_mm_dd_yyyy)
					if report_success == True:
						message = "Report Successfully Appended" + message
					elif report_success == False:
						message = "Report Fail - Please save all open word documents and try again."
				else:
					message = "%s%s" % (chart_message, message)
					
				#Log the success/failure of the report generation and the query conditions to the event logger
				self.EventLogger(message)
			
			#Open Word document for review by the user
			self.close_word_document(False)
		else:
			#If any of the query conditions are invalid, most of the failure modes 
			#are related to the selected (or not selected) report *.docx file
			#see "Initial user input checks" at the beginning of this method for a complete of 
			#pre-run checks
			message = "Report Fail - Due to the listed reason(s) no report was generated.\n\tPlease make the necessary corrections to your query\n\tconditions then resubmit the report request."
			self.EventLogger(message)
	
		self.EventLogger('Action Complete\n')
	
	def get_averaging_denominator(self, average_by_index, start_date_yyyy_mm_dd, end_date_yyyy_mm_dd):
		#average_by_index
			#0 = Day
			#1 = Week
			#2 = Month
			#3 = Quarter
			#4 = Year
		
		day_count = self.calculate_day_delta(start_date_yyyy_mm_dd, end_date_yyyy_mm_dd)
		days_in_month = 365.00 / 12.00
		
		if average_by_index == 0:
			denominator = day_count
		elif average_by_index == 1:
			denominator = day_count/7.00
		elif average_by_index == 2:
			denominator = day_count/days_in_month
		elif average_by_index == 3:
			denominator = day_count/(days_in_month * 3)
		elif average_by_index == 4:
			denominator = day_count/365.00
			
		return denominator
		
	def calculate_day_delta(self, start_date, end_date):
		#Returns the number of days in the defined date range
		#NOT the difference between the two dates.  
		#If the start_date and end_date are the same it
		#will return a value of 1
		start_day = datetime.datetime.strptime(start_date, "'%Y-%m-%d'")
		end_day = datetime.datetime.strptime(end_date, "'%Y-%m-%d'")
		delta = end_day - start_day
		return delta.days + 1
	
	def close_word_document(self, action):
		if action == True:
			word = win32.gencache.EnsureDispatch('Word.Application')
			word.ActiveDocument.SaveAs(self.ReportFilePath)
			word.Quit()
			
		if action == False:
			word = win32.gencache.EnsureDispatch('Word.Application')
			word.Visible = True
			word.Documents.Open(self.ReportFilePath)
			
	
	def AppendDocument(self, new_report, data_label_list, data_value_list, population_selection, percentages, category_selection, start_date_mm_dd_yyyy, end_date_mm_dd_yyyy):
	
		if new_report == True:
			report_doc = Document(self.WordTemplate_file_path)
		else:
			report_doc = Document(self.ReportFilePath)
		
		report_doc.add_page_break()
		
		section = report_doc.sections
		
		section[0].top_margin = Inches(0.5)
		section[0].bottom_margin = Inches(0.5)
		section[0].left_margin = Inches(0.75)
		section[0].right_margin = Inches(0.75)
		
		if population_selection == 'All':
			data_heading_lvl_1 = "All Employees\tCategory: %s" % (category_selection)
		else:
			data_heading_lvl_1 = "%s\t\tCategory: %s" % (population_selection, category_selection)
			
		report_doc.add_heading(data_heading_lvl_1, level=1)
		
		bullet_text = []
		bullet_text.append("Department: %s" % self.DepartmentSelected)
		bullet_text.append("Data Aggregation: %s%s" % (self.hour_label, self.average_by_label))
		bullet_text.append("%s  -  %s" % (start_date_mm_dd_yyyy, end_date_mm_dd_yyyy))
		
		for text_line in bullet_text:
			paragraph = report_doc.add_paragraph(text_line)
			paragraph.style = 'List Bullet'
		
		'''jump'''
		if sum(data_value_list) != 0:
			report_doc.add_picture('temp_bar_chart.png', width=Inches(7))
		else:
			
			first_line = "This query returned a sum of zero hours."
			second_line = "\nThe following lists possibilities why this may have occurred:"
			paragraph = report_doc.add_paragraph("\n\n\n")
			paragraph.add_run(first_line).bold = True
			paragraph.add_run(second_line).bold = True
			
			bullet_text = []
			bullet_text.append("No activity sheets have been imported into the database for the selected employee population within the defined date range.")
			bullet_text.append("If you are certain you have imported activity sheet(s) matching your query, check the date(s) on the activity sheet(s), it may be inaccurate.")
			bullet_text.append("The employee population worked zero total hours for the queried categories within the defined date range.")
			bullet_text.append("This is not an exhaustive list, just the most common reasons.  There are other possibilities for a zero sum query.")
			
			for text_line in bullet_text:
				paragraph = report_doc.add_paragraph(text_line)
				paragraph.style = 'List Bullet'
			
			paragraph = report_doc.add_paragraph("\n\n\n")

		os.remove("%stemp_bar_chart.png" % self.ActiveDirectory)
		
		table = report_doc.add_table(rows=1, cols=3)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Category'
		hdr_cells[1].text = "%s%s" % (self.hour_label, self.average_by_label)
		hdr_cells[2].text = '%'

		for x in range(len(data_label_list)):
			row_cells = table.add_row().cells
			row_cells[0].text = str(data_label_list[x])
			row_cells[1].text = str(data_value_list[x])
			percentages[x] = round(percentages[x],1)
			row_cells[2].text = str(percentages[x])
		
		row_cells = table.add_row().cells
		row_cells[0].text = 'Total'
		row_cells[1].text = str(sum(data_value_list))
		row_cells[2].text = '100'
		table.style = 'Grid Table 4'
		
		try:
			self.ReportFilePath
			report_doc.save(self.ReportFilePath)
			report_success = True
		except IOError:
			report_success = False
			
		return report_success
	
	def AddDays(self, start_date, days_to_add):

		date_1 = datetime.datetime.strptime(start_date, '%Y-%m-%d')
		end_date = date_1 + datetime.timedelta(days=days_to_add)
		return end_date.strftime('%Y-%m-%d')
		
	def CreateBarChart(self, data_label_list, data_value_list, columns_list):
		figure1 = DataVisualizer()
				
		data_label_list.reverse()
		data_value_list.reverse()
		
		data_labels = data_label_list
		data_values = data_value_list
		x_axis_label = "%s%s" % (self.hour_label, self.average_by_label)
		chart_title = ''
		
		if len(columns_list) == 1:
			bar_color = 'b'
		else:
			bar_color = 'g'
		
		figure1.create_horizontal_bar_chart(data_labels, data_values, bar_color, x_axis_label, chart_title)
		
		data_label_list.reverse()
		data_value_list.reverse()
		data_labels = data_label_list
		data_values = data_value_list
		
		if sum(data_value_list) == 0:
			percentages = [x for x in data_value_list]
		else:
			percentages = figure1.calc_percentages(data_value_list)
		
		message = "Chart Successfully Created"
		chart_created = True
		
		return percentages, chart_created, message
		
	
	def OnUpdateEmployee(self, event):

		#Get employee information from GUI widgets
		clk_no = self.clock_no_txtCtrl.GetValue()
		emp_name = self.employee_name_txtCtrl.GetValue()
		
		#Verify that the clock number & employee name fields are not empty
		if not(clk_no == "" or emp_name == "") and len(clk_no) == 4:
		
			lst_index = self.type_listbox.GetSelection()
			emp_type = str("%s" % self.type_listbox.GetString(lst_index))	

			#message = "Name: %s Clk No:  %s Type: %s" % (emp_name, clk_no, emp_type)
			
			#self.EventLogger(message)
			
			#Query clock number to check if it is already in the database
			cursor = self.empDB.conn.execute("SELECT employee_name FROM employee WHERE clk_no = ?", (clk_no,))
			query_emp_name = cursor.fetchall()
			
			if len(query_emp_name) == 0:
				#Employee clock number not found
				#So prompt if a new employee should be added to the database
				#with the defined information
				new_employee = True
				question = '''Do you wish to add the following employee to the database?
				
								Name: %s
								Clock Number: %s
								Type: %s''' % (emp_name, clk_no, emp_type)
				caption = "ADD NEW EMPLOYEE"
				
			else:
				#Employee clock number found
				#so prompt if the selected employee should be updated
				#with the defined information
				new_employee = False
				question = '''Do you wish to update the selected employee already
								in the database with the defined information below?
					
								Name: %s
								Clock Number: %s
								Type: %s''' % (emp_name, clk_no, emp_type)
				caption = "CHANGE EXISTING EMPLOYEE INFORMATION"
				
			#Yes/No dialogue box
			dlg = wx.MessageDialog(self, question, caption, wx.YES_NO | wx.ICON_QUESTION)
			make_db_changes = dlg.ShowModal() == wx.ID_YES
			dlg.Destroy()
			
			if make_db_changes == True and new_employee == True:
				#Add new employee to the database
				table = 'employee'
				column = ['clk_no', 'employee_name', 'employee_type']
				values = [str(clk_no), str(emp_name), str(emp_type)]
				self.empDB.Insert_Query_No_Conditions(table, column, values)
				
				message = 'New Employee added to the database:\n\tName: %s\n\tClock No: %s\n\tType: %s' % (emp_name, clk_no, emp_type)
			elif make_db_changes == True and new_employee == False:
				#Make the defined changes to the existing employee
				
				if emp_name == "delete":
					#Back door method of deleting an employee record
					self.empDB.conn.execute("DELETE FROM employee WHERE clk_no = ?", (clk_no,))
					self.empDB.conn.commit()
					message = 'Employee deleted from the database:\n\tName: %s\n\tClock No: %s\n\tType: %s' % (emp_name, clk_no, emp_type)
				else:
					#Edit an existing employee
					self.empDB.conn.execute("UPDATE employee SET employee_name = ?, employee_type = ? WHERE clk_no = ?", (emp_name, emp_type, clk_no,))
					self.empDB.conn.commit()
					
					message = 'Employee information updated:\n\tName: %s\n\tClock No: %s\n\tType: %s' % (emp_name, clk_no, emp_type)
					
			elif make_db_changes == False:
				#Do not make any changes to the database
				message = "No Changes made to the employee database"
				
			self.Update_Population_ListBox()
			
		else:
			#Invalid employee information provided feed back
			if clk_no == "":
				message = "Clock number field is blank. Cannot edit employee information."
			elif emp_name == "":
				message = "Employee Name field is blank. Cannot edit employee information."
			elif not(len(clk_no) == 4):
				message = "Clock number is not 4 characters in length. Cannot edit employee information."

		self.EventLogger(message)
		self.EventLogger('Action Complete\n')
		
	def OnEmployeeInfo(self, event):
		lst_index = self.emp_listbox.GetSelection()
		
		message = 'Employee Information Query\n'
		self.EventLogger(message)
		
		# If the selection is an individual employee
		if lst_index > 3:
			#Get employee select number & name for corresponding GUI list box
			
			emp_name = str("%s" % self.emp_listbox.GetString(lst_index))
			emp_info_lst = [emp_name]
			
			#SELECT query employee table for clock number & type
			cursor = self.empDB.conn.execute("SELECT clk_no, employee_type FROM employee WHERE employee_name = ?", (emp_name,))
			emp_info = cursor.fetchall()
			
			#Populate queried employee information into emp_info_list
			for info in emp_info[0]:
				emp_info_lst.append(str(info))
				
			#SELECT query upload_dates table for list of mondays_date for given employee
			cursor = self.empDB.conn.execute("SELECT mondays_date FROM upload_dates WHERE employee_clk_no = ?", (emp_info_lst[1],))
			emp_dates = cursor.fetchall()
			
			#Populate queried employee dates into emp_dates_list
			emp_dates_list = []
			for i ,d in enumerate(emp_dates):
				emp_dates_list.append(str("%s" % d))
				emp_dates_list[i] = str(datetime.datetime.strptime(emp_dates_list[i], '%Y-%m-%d').strftime('%m/%d/%Y'))
				
			message = 'Name:  %s\t  Clock No:  %s\tType:  %s\nImported Activity Sheets:' % (emp_info_lst[0], emp_info_lst[1], emp_info_lst[2])
			self.status_logger.AppendText(message)
			for d in emp_dates_list:
				message = '\n\t%s' % str(d)
				self.status_logger.AppendText(message)
		
		#If the selection represents more than a single employee
		#the query will return information on all employees as well as
		#the most recent upload date
		else:
			#get clock numbers of all the employees in the database
			cursor = self.empDB.conn.execute("SELECT clk_no FROM employee ORDER BY employee_name ASC")
			clock_numbers = cursor.fetchall()
			message = "%s        \t%s  \t%s  \t      %s\n" % ("Name", "No.", "Type", "Most Recent Upload")
			self.status_logger.AppendText(message)
			for cn in clock_numbers:
				clock_number = "%s" % cn
				
				cursor = self.empDB.conn.execute("SELECT employee.employee_name, employee.clk_no, employee.employee_type, \
														MAX(upload_dates.mondays_date) \
												FROM employee \
												JOIN upload_dates ON \
													employee.clk_no = upload_dates.employee_clk_no	\
												WHERE \
													employee.clk_no = ?", (clock_number,))
												
				message = cursor.fetchall()
				for row in message:
					
					#if there are no dates in the database for a given clock number
					if row[3] == None:
						#print clock_number
						#print "hello"
						message = self.empDB.conn.execute("SELECT employee_name, clk_no, employee_type FROM employee \
												WHERE clk_no = ?", (clock_number,))
						for row in message:
							message = "%s  \t%s  \t%s  \t      None\n" % (row[0], row[1], row[2])
					
					#if there are dates in the database for a given clock number
					else:
					
						if len(row[0]) > 13:
							truncated_name = row[0][:13]
						else:
							truncated_name = row[0]
							
						if row[2] == 'In-House':
							truncated_type = 'St Joe'
						else:
							truncated_type = row[2]
						
						report_date = str(datetime.datetime.strptime(row[3], '%Y-%m-%d').strftime('%m/%d/%Y'))
						
						if self.real_connect == True:
							message = "%s \t%s  \t%s    \t%s\n" % (truncated_name, row[1], truncated_type, report_date)
						else:
							message = "%s\t%s  \t%s  \t%s\n" % (truncated_name, row[1], truncated_type, report_date)
						
					self.status_logger.AppendText(message)
		
		self.status_logger.AppendText("\n")
		self.EventLogger('Action Complete\n')
		
	
	def EvtOutputOptionsRadioBox(self, event):
		result = event.GetInt()
		if result == 0:
			self.ave_by_listbox.Disable()
		elif result == 1:
			self.ave_by_listbox.Enable()
			
	def EvtDepartmentRadioBox(self, event):
		result = event.GetInt()
		if result == 0:
			self.Update_Category_ListBox("SepSci")
			self.DepartmentSelected = "SepSci"
		elif result == 1:
			self.Update_Category_ListBox("GDS")
			self.DepartmentSelected = "GDS"
		
	def EventLogger(self, message):
		current_time = str(time.strftime("%H:%M:%S"))
		current_date = str(time.strftime("%m/%d/%Y"))
		status_message = "%s %s - %s\n" % (current_date, current_time, message) 
		self.status_logger.AppendText(status_message)
		
	def Get_Data_Dates(self):
	
		#Get all the distinct dates from the catDB.
		#Since each table in the catDB has the same
		#data_dates the Miscellaneous table was 
		#arbitrarily selected to query
		table = 'Miscellaneous'
		columns = 'DISTINCT data_date'
		conditions = None
		data_date_lst = []
		cursor = self.catDB.Select_Query(columns, table, conditions)
		for index, row in enumerate(cursor):
			data_date_lst.append(str(datetime.datetime.strptime(str('%s' % row), '%Y-%m-%d').strftime('%m/%d/%Y  (%a)')))
		if data_date_lst == []:
			data_date_lst = ['No Data']
			
		return data_date_lst
		
	def Update_Date_ListBoxes(self):
		data_date_lst = self.Get_Data_Dates()
		
		if data_date_lst[0] != 'No Data':
			last_date = len(data_date_lst)-1
			last_date_minus_a_week = len(data_date_lst)-7
		else:
			last_date = 0
			last_date_minus_a_week = 0
		
		self.start_date_listbox.Clear()
		self.end_date_listbox.Clear()
		for d in data_date_lst:
			self.start_date_listbox.Append(d)
			self.end_date_listbox.Append(d)
			
		self.start_date_listbox.SetSelection(last_date_minus_a_week)
		self.end_date_listbox.SetSelection(last_date)
		
	def Update_Category_ListBox(self, LECO_group):
		self.cat_listbox.Clear()
		cat_list = []
		if LECO_group == "SepSci":
			cat_list = self.display_sep_sci_cat_list
		elif LECO_group == "GDS":
			cat_list = self.display_gds_cat_list
		
		for c in cat_list:
			self.cat_listbox.Append(c)
	
		self.cat_listbox.SetSelection(0)
		
	def Update_Population_ListBox(self):
		self.emp_listbox.Clear()
		#Get the employee names from the employee tables
		#and put them into a list.  Then insert into the list
		#'all', 'field service only', 'in-house only'
		emp_names_DB_Object = self.empDB.Select_Query('employee_name', 'employee', 'ORDER BY employee_name ASC')
		self.emp_names_list = []
		for name_tuple in emp_names_DB_Object:
			name_str = "%s" % name_tuple
			self.emp_names_list.append(str(name_str))
		self.emp_names_list.insert(0,'Field Service')
		self.emp_names_list.insert(0,'In-House')
		self.emp_names_list.insert(0,'All')
		
		for p in self.emp_names_list:
			self.emp_listbox.Append(p)
	
		self.emp_listbox.SetSelection(0)
	
app = wx.App(False)
frame = ManHourTrackerFrame(None, title="TIME TRAX 2.2: Man Hour Tracker")
panel = ManHourTrackerPanel(frame)
frame.Show(True)
app.MainLoop()
